"""
📚 Писалка-читалка - PYQT6 EDITION
Версия: 1.2
"""

import os
import re
import sys
import xml.etree.ElementTree as ET
import xml.dom.minidom
from datetime import datetime
from typing import Optional, Dict, List
from PyQt6.QtWidgets import *
from PyQt6.QtCore import *
from PyQt6.QtGui import *
from PyQt6.QtPrintSupport import QPrinter
from PyQt6.QtCore import pyqtSignal
from PyQt6.QtGui import QDesktopServices
from PyQt6.QtCore import QUrl

from config import Theme, Settings, Config
from data_model import Project, Chapter
from utils import resource_path, ICON_PATH
from updater import UpdateChecker
from search import SearchEngine

# === ИМПОРТ ЭКСПОРТЕРОВ ===
try:
    from exporters.fb2_exporter import FB2Exporter
except ImportError:
    FB2Exporter = None

try:
    from exporters.docx_exporter import DOCXExporter
except ImportError:
    DOCXExporter = None

try:
    from exporters.pdf_exporter import PDFExporter
except ImportError:
    PDFExporter = None

try:
    from exporters.txt_exporter import TXTExporter
except ImportError:
    TXTExporter = None

# === ИМПОРТ ДЛЯ PDF ===
try:
    from PyQt6.QtPdf import QPdfDocument
    from PyQt6.QtPdfWidgets import QPdfView

    PDF_SUPPORT = True
except ImportError:
    QPdfDocument = None
    QPdfView = None
    PDF_SUPPORT = False

try:
    import PyPDF2

    PYPDF2_SUPPORT = True
except ImportError:
    PyPDF2 = None
    PYPDF2_SUPPORT = False


class BookViewer(QDialog):
    """Диалог для предпросмотра книг как в FBReader"""

    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("📖 Читалка книг")
        self.setGeometry(200, 200, 900, 700)

        # Настройка внешнего вида как у FBReader
        self.setStyleSheet("""
            QDialog {
                background-color: #f5f0e8;
            }
            QTextEdit {
                background-color: #fbf8f0;
                color: #2c2c2c;
                font-family: 'Georgia', serif;
                font-size: 14px;
                border: none;
                padding: 20px;
            }
            QToolBar {
                background-color: #e8e0d8;
                border: none;
                spacing: 5px;
            }
            QToolButton {
                background-color: transparent;
                border: none;
                padding: 5px;
                font-size: 16px;
            }
            QToolButton:hover {
                background-color: #d0c8c0;
                border-radius: 3px;
            }
        """)

        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        # Панель инструментов как у FBReader
        toolbar = QToolBar()
        toolbar.setMovable(False)
        toolbar.setIconSize(QSize(24, 24))

        # Кнопки навигации
        self.back_btn = QAction("◀", self)
        self.back_btn.triggered.connect(self.prev_page)
        toolbar.addAction(self.back_btn)

        self.forward_btn = QAction("▶", self)
        self.forward_btn.triggered.connect(self.next_page)
        toolbar.addAction(self.forward_btn)

        toolbar.addSeparator()

        # Кнопки масштаба
        zoom_out = QAction("🔍−", self)
        zoom_out.triggered.connect(lambda: self.change_font_size(-1))
        toolbar.addAction(zoom_out)

        zoom_in = QAction("🔍+", self)
        zoom_in.triggered.connect(lambda: self.change_font_size(1))
        toolbar.addAction(zoom_in)

        toolbar.addSeparator()

        # Кнопки открытия
        open_action = QAction("📂 Открыть", self)
        open_action.triggered.connect(self.open_book)
        toolbar.addAction(open_action)

        # Поиск
        search_action = QAction("🔍 Поиск", self)
        search_action.triggered.connect(self.find_text)
        toolbar.addAction(search_action)

        # Информация о книге
        self.info_label = QLabel("Нет книги")
        self.info_label.setStyleSheet("padding: 5px; color: #555; font-size: 12px;")
        toolbar.addWidget(self.info_label)

        toolbar.addSeparator()

        # Кнопка полного экрана
        fullscreen_action = QAction("⛶", self)
        fullscreen_action.triggered.connect(self.toggle_fullscreen)
        toolbar.addAction(fullscreen_action)

        layout.addWidget(toolbar)

        # Основной виджет для текста
        self.text_viewer = QTextEdit()
        self.text_viewer.setReadOnly(True)
        self.text_viewer.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)

        # Настройка шрифта как в FBReader
        font = QFont("Georgia", 14)
        self.text_viewer.setFont(font)
        self.font_size = 14

        layout.addWidget(self.text_viewer)

        # Статусбар
        self.status_label = QLabel()
        self.status_label.setStyleSheet("padding: 5px; background-color: #e8e0d8; color: #555;")
        layout.addWidget(self.status_label)

        self.setLayout(layout)

        self.current_file = None
        self.current_page = 0
        self.total_pages = 0
        self.chapters = []
        self.fullscreen = False
        self.pdf_view = None
        self.current_text = ""

    def toggle_fullscreen(self):
        """Переключить полноэкранный режим"""
        if self.fullscreen:
            self.showNormal()
        else:
            self.showFullScreen()
        self.fullscreen = not self.fullscreen

    def change_font_size(self, delta: int):
        """Изменить размер шрифта"""
        self.font_size = max(8, min(30, self.font_size + delta))
        self.text_viewer.setFont(QFont("Georgia", self.font_size))
        self.status_label.setText(f"Размер шрифта: {self.font_size}pt")

    def prev_page(self):
        """Предыдущая страница"""
        scrollbar = self.text_viewer.verticalScrollBar()
        scrollbar.setValue(scrollbar.value() - scrollbar.pageStep())

    def next_page(self):
        """Следующая страница"""
        scrollbar = self.text_viewer.verticalScrollBar()
        scrollbar.setValue(scrollbar.value() + scrollbar.pageStep())

    def find_text(self):
        """Поиск текста в книге"""
        if not self.current_text and not self.text_viewer.toPlainText():
            QMessageBox.information(self, "Поиск", "Сначала откройте книгу")
            return

        text, ok = QInputDialog.getText(
            self,
            "Поиск",
            "Введите текст для поиска:",
            QLineEdit.EchoMode.Normal
        )

        if ok and text:
            content = self.text_viewer.toPlainText()
            if text.lower() in content.lower():
                pos = content.lower().find(text.lower())
                cursor = self.text_viewer.textCursor()
                cursor.setPosition(pos)
                cursor.setPosition(pos + len(text), QTextCursor.MoveMode.KeepAnchor)
                self.text_viewer.setTextCursor(cursor)
                self.text_viewer.setFocus()
                self.status_label.setText(f"Найдено: '{text}'")
            else:
                QMessageBox.information(self, "Поиск", "Текст не найден")

    def open_book(self):
        """Открыть файл книги"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите книгу для чтения",
            os.path.expanduser("~"),
            "Книги (*.txt *.fb2 *.pdf *.docx)"
        )

        if file_path:
            self.load_book(file_path)

    def load_book(self, file_path: str):
        """Загрузить книгу с нормальным отображением"""
        try:
            ext = os.path.splitext(file_path)[1].lower()
            self.current_file = file_path
            self.chapters = []
            content = ""

            # Если есть PDF виджет, удаляем его
            if self.pdf_view:
                self.pdf_view.deleteLater()
                self.pdf_view = None
                # Возвращаем текстовый виджет
                if self.text_viewer not in self.children():
                    layout = self.layout()
                    layout.replaceWidget(self.pdf_view, self.text_viewer)

            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.current_text = content
                self.text_viewer.setPlainText(content)
                word_count = len(re.findall(r'\b\w+\b', content))
                self.status_label.setText(f"📄 TXT • {word_count} слов • {len(content)} символов")

            elif ext == '.fb2':
                content = self.parse_fb2(file_path)
                self.current_text = content
                self.text_viewer.setHtml(content)
                self.status_label.setText("📚 FB2 • Загружено")

            elif ext == '.pdf':
                if PDF_SUPPORT:
                    # Используем Qt PDF
                    self.pdf_view = QPdfView()
                    pdf_doc = QPdfDocument()
                    pdf_doc.load(file_path)
                    self.pdf_view.setDocument(pdf_doc)

                    # Заменяем виджет
                    self.text_viewer.hide()
                    layout = self.layout()
                    layout.replaceWidget(self.text_viewer, self.pdf_view)
                    self.pdf_view.show()
                    self.status_label.setText(f"📕 PDF • {os.path.basename(file_path)}")
                elif PYPDF2_SUPPORT:
                    # Fallback на PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text_parts = []
                        for page in reader.pages:
                            text_parts.append(page.extract_text())
                        content = "\n\n".join(text_parts)
                    self.current_text = content
                    self.text_viewer.setPlainText(content)
                    self.status_label.setText(f"📕 PDF (текст) • {len(content)} символов")
                else:
                    content = f"<h1>📕 PDF файл</h1>"
                    content += f"<p><b>{os.path.basename(file_path)}</b></p>"
                    content += "<p style='color: #666;'>Установите PyQt6-Pdf или PyPDF2 для просмотра PDF</p>"
                    self.text_viewer.setHtml(content)
                    self.status_label.setText("⚠️ PDF поддержка ограничена")

            elif ext == '.docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    html = self.parse_docx(doc)
                    self.current_text = html
                    self.text_viewer.setHtml(html)
                    self.status_label.setText(f"📝 DOCX • Загружено")
                except ImportError:
                    content = f"<h1>📝 DOCX файл</h1>"
                    content += f"<p><b>{os.path.basename(file_path)}</b></p>"
                    content += "<p style='color: #666;'>Установите python-docx для просмотра DOCX</p>"
                    self.text_viewer.setHtml(content)
                    self.status_label.setText("⚠️ DOCX поддержка ограничена")

            # Обновляем заголовок
            self.setWindowTitle(f"📖 {os.path.basename(file_path)} - Читалка книг")
            self.info_label.setText(os.path.basename(file_path))

        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить книгу:\n{e}")

    def parse_fb2(self, file_path: str) -> str:
        """Парсинг FB2 в красивый HTML"""
        try:
            # Убираем неймспейс для простоты
            it = ET.iterparse(file_path)
            for _, el in it:
                if '}' in el.tag:
                    el.tag = el.tag.split('}', 1)[1]
            root = it.root

            html = "<html><body style='max-width: 800px; margin: 0 auto; padding: 20px; font-family: Georgia, serif;'>"

            # Информация о книге
            title_info = root.find(".//title-info")
            if title_info is not None:
                book_title = title_info.find("book-title")
                if book_title is not None and book_title.text:
                    html += f"<h1 style='text-align: center; color: #2c2c2c;'>{book_title.text}</h1>"

                # Автор
                author = title_info.find("author")
                if author is not None:
                    first = author.find("first-name")
                    last = author.find("last-name")
                    if first is not None and first.text and last is not None and last.text:
                        html += f"<p style='text-align: center; color: #666;'><i>{first.text} {last.text}</i></p>"

                # Аннотация
                annotation = title_info.find("annotation")
                if annotation is not None:
                    html += "<div style='background: #f5f0e8; padding: 15px; margin: 20px 0; border-left: 4px solid #8b6b4f;'>"
                    for p in annotation.findall("p"):
                        if p.text:
                            html += f"<p style='font-style: italic;'>{p.text}</p>"
                    html += "</div>"

            # Тело книги
            body = root.find("body")
            if body is not None:
                for section in body.findall("section"):
                    # Название главы
                    title = section.find("title")
                    if title is not None:
                        p_title = title.find("p")
                        if p_title is not None and p_title.text:
                            html += f"<h2 style='margin-top: 40px; color: #4a3a2a;'>{p_title.text}</h2>"

                    # Содержимое главы
                    for p in section.findall("p"):
                        if p.text:
                            html += f"<p style='line-height: 1.8; text-align: justify;'>{p.text}</p>"
                        elif p.find("empty-line") is not None:
                            html += "<br>"

            html += "</body></html>"
            return html

        except Exception as e:
            return f"<p><b>Ошибка парсинга FB2:</b> {e}</p>"

    def parse_docx(self, doc) -> str:
        """Парсинг DOCX в HTML"""
        html = "<html><body style='max-width: 800px; margin: 0 auto; padding: 20px; font-family: Georgia, serif;'>"

        for para in doc.paragraphs:
            if para.text.strip():
                # Проверяем стиль параграфа
                style_name = para.style.name if para.style else ""
                if 'Heading' in style_name:
                    level = 1
                    for i in range(1, 6):
                        if f"Heading {i}" in style_name:
                            level = i
                            break
                    html += f"<h{level}>{para.text}</h{level}>"
                else:
                    html += f"<p style='line-height: 1.8; text-align: justify;'>{para.text}</p>"

        html += "</body></html>"
        return html


class CoverDialog(QDialog):
    """Диалог для управления обложкой"""

    def __init__(self, project: Project, parent=None):
        super().__init__(parent)
        self.project = project
        self.setWindowTitle("Обложка книги")
        self.setMinimumWidth(500)
        self.setMinimumHeight(400)

        layout = QVBoxLayout()

        # Заголовок
        title_label = QLabel(f"Проект: {project.name}")
        title_label.setStyleSheet("font-size: 16px; font-weight: bold; margin: 10px;")
        layout.addWidget(title_label)

        # Область отображения обложки
        self.cover_label = QLabel()
        self.cover_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.cover_label.setMinimumHeight(300)
        self.cover_label.setStyleSheet("border: 2px dashed gray; margin: 10px;")

        # Загружаем текущую обложку
        self.update_cover_display()

        layout.addWidget(self.cover_label)

        # Информация о текущей обложке
        self.info_label = QLabel()
        self.update_info_label()
        layout.addWidget(self.info_label)

        # Кнопки
        btn_layout = QHBoxLayout()

        self.add_btn = QPushButton("📷 Добавить обложку")
        self.add_btn.clicked.connect(self.add_cover)

        self.remove_btn = QPushButton("🗑 Удалить обложку")
        self.remove_btn.clicked.connect(self.remove_cover)
        self.remove_btn.setEnabled(self.project.cover_path is not None)

        btn_layout.addWidget(self.add_btn)
        btn_layout.addWidget(self.remove_btn)

        layout.addLayout(btn_layout)

        # Кнопки диалога
        dialog_buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok |
            QDialogButtonBox.StandardButton.Cancel
        )
        dialog_buttons.accepted.connect(self.accept)
        dialog_buttons.rejected.connect(self.reject)
        layout.addWidget(dialog_buttons)

        self.setLayout(layout)

    def update_cover_display(self):
        """Обновить отображение обложки"""
        pixmap = self.project.get_cover_pixmap(300, 400)
        if pixmap:
            self.cover_label.setPixmap(pixmap)
        else:
            self.cover_label.setText("📖\n\nНет обложки\n\nНажмите 'Добавить обложку' для загрузки")
            self.cover_label.setStyleSheet("border: 2px dashed gray; margin: 10px; font-size: 16px;")

    def update_info_label(self):
        """Обновить информацию об обложке"""
        if self.project.cover_path and os.path.exists(self.project.cover_path):
            size = os.path.getsize(self.project.cover_path)
            name = os.path.basename(self.project.cover_path)
            self.info_label.setText(f"Текущая обложка: {name} ({size:,} байт)")
            self.info_label.setStyleSheet("color: green; margin: 5px;")
        else:
            self.info_label.setText("Обложка не установлена")
            self.info_label.setStyleSheet("color: gray; margin: 5px;")

    def add_cover(self):
        """Добавить обложку"""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите изображение для обложки",
            os.path.expanduser("~"),
            "Изображения (*.png *.jpg *.jpeg *.bmp *.gif)"
        )

        if file_path:
            if self.project.set_cover(file_path):
                self.project.save()
                self.update_cover_display()
                self.update_info_label()
                self.remove_btn.setEnabled(True)
                QMessageBox.information(self, "Успех", "Обложка успешно добавлена!")
            else:
                QMessageBox.warning(self, "Ошибка", "Не удалось установить обложку. Проверьте формат файла.")

    def remove_cover(self):
        """Удалить обложку"""
        reply = QMessageBox.question(
            self,
            "Подтверждение",
            "Удалить обложку?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.project.remove_cover()
            self.project.save()
            self.update_cover_display()
            self.update_info_label()
            self.remove_btn.setEnabled(False)
            QMessageBox.information(self, "Успех", "Обложка удалена")


class ChapterTreeItem(QTreeWidgetItem):
    """Элемент дерева глав"""

    def __init__(self, chapter: Chapter):
        super().__init__()
        self.chapter = chapter
        self.update_display()

    def update_display(self):
        """Обновить отображение"""
        # Используем эмодзи вместо иконок (работает на всех платформах)
        status_icons = {
            'done': '✅',
            'review': '🔍',
            'draft': '📝'
        }
        icon = status_icons.get(self.chapter.status, '📝')
        self.setText(0, f"{icon} {self.chapter.order + 1}. {self.chapter.title}")
        self.setToolTip(0, f"Слов: {self.chapter.word_count()}\nСтатус: {self.chapter.status}")


class ChapterTitleEditor(QWidget):
    """Виджет для редактирования названия главы"""
    title_changed = pyqtSignal(str)

    def __init__(self, chapter: Chapter, parent=None):
        super().__init__(parent)
        self.chapter = chapter

        layout = QHBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)

        self.title_edit = QLineEdit(chapter.title)
        self.title_edit.editingFinished.connect(self.on_title_changed)

        layout.addWidget(self.title_edit)
        self.setLayout(layout)

    def on_title_changed(self):
        """Когда название изменено"""
        new_title = self.title_edit.text().strip()
        if new_title and new_title != self.chapter.title:
            self.title_changed.emit(new_title)


class CodeEditor(QPlainTextEdit):
    """Редактор кода с подсветкой и Vim-режимом"""

    def __init__(self, parent=None):
        super().__init__(parent)

        font = QFont('Consolas', 12)
        if not QFontInfo(font).fixedPitch():
            font = QFont('Courier New', 12)
        self.setFont(font)

        self.setLineWrapMode(QPlainTextEdit.LineWrapMode.WidgetWidth)

        self.tab_size = 4
        self.vim_mode = True
        self.mode = 'insert'

        self.status_label = QLabel("-- ВСТАВКА --")
        self.status_label.setStyleSheet("color: gray; padding: 2px;")

        self.highlight_current_line()

    def highlight_current_line(self):
        """Подсветка текущей строки"""
        try:
            extra_selections = []

            if not self.isReadOnly():
                selection = QTextEdit.ExtraSelection()
                line_color = QColor(Qt.GlobalColor.yellow).lighter(180)
                selection.format.setBackground(line_color)
                selection.format.setProperty(QTextFormat.Property.FullWidthSelection, True)
                selection.cursor = self.textCursor()
                selection.cursor.clearSelection()
                extra_selections.append(selection)

            self.setExtraSelections(extra_selections)
        except Exception as e:
            print(f"[CodeEditor] Ошибка подсветки: {e}")

    def keyPressEvent(self, event):
        """Обработка клавиш"""
        try:
            if not self.vim_mode:
                super().keyPressEvent(event)
                return

            if self.mode == 'normal':
                key = event.key()

                if key == Qt.Key.Key_I:
                    self.mode = 'insert'
                    self.status_label.setText("-- ВСТАВКА --")
                elif key == Qt.Key.Key_H:
                    self.moveCursor(QTextCursor.MoveOperation.Left)
                elif key == Qt.Key.Key_J:
                    self.moveCursor(QTextCursor.MoveOperation.Down)
                elif key == Qt.Key.Key_K:
                    self.moveCursor(QTextCursor.MoveOperation.Up)
                elif key == Qt.Key.Key_L:
                    self.moveCursor(QTextCursor.MoveOperation.Right)
                elif key == Qt.Key.Key_X:
                    self.textCursor().deleteChar()
                elif key == Qt.Key.Key_U and event.modifiers() & Qt.KeyboardModifier.ControlModifier:
                    self.undo()

            elif self.mode == 'insert':
                if event.key() == Qt.Key.Key_Escape:
                    self.mode = 'normal'
                    self.status_label.setText("-- НОРМАЛЬНЫЙ --")
                else:
                    super().keyPressEvent(event)

            self.highlight_current_line()
        except Exception as e:
            print(f"[CodeEditor] Ошибка в keyPressEvent: {e}")
            super().keyPressEvent(event)


class SettingsDialog(QDialog):
    """Диалог настроек приложения"""
    theme_changed = pyqtSignal(str)

    def __init__(self, settings: Settings, parent=None):
        super().__init__(parent)
        self.settings = settings
        self.setWindowTitle("Настройки")
        self.setMinimumWidth(600)
        self.setMinimumHeight(500)

        layout = QVBoxLayout()

        # Создаем вкладки
        tabs = QTabWidget()

        # Вкладка "Общие"
        general_tab = self.create_general_tab()
        tabs.addTab(general_tab, "Общие")

        # Вкладка "Темы"
        theme_tab = self.create_theme_tab()
        tabs.addTab(theme_tab, "Оформление")

        # Вкладка "Редактор"
        editor_tab = self.create_editor_tab()
        tabs.addTab(editor_tab, "Редактор")

        layout.addWidget(tabs)

        # Кнопки
        btn_layout = QHBoxLayout()
        save_btn = QPushButton("Сохранить")
        save_btn.clicked.connect(self.save_settings)
        cancel_btn = QPushButton("Отмена")
        cancel_btn.clicked.connect(self.reject)
        apply_btn = QPushButton("Применить")
        apply_btn.clicked.connect(self.apply_settings)

        btn_layout.addWidget(save_btn)
        btn_layout.addWidget(apply_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

        self.setLayout(layout)

    def create_general_tab(self) -> QWidget:
        """Создать вкладку общих настроек"""
        widget = QWidget()
        layout = QVBoxLayout()

        # Автосохранение
        auto_save_group = QGroupBox("Автосохранение")
        auto_save_layout = QVBoxLayout()

        self.auto_save_check = QCheckBox("Включить автосохранение")
        self.auto_save_check.setChecked(self.settings.config.get('auto_save_enabled', True))

        auto_save_interval_layout = QHBoxLayout()
        auto_save_interval_layout.addWidget(QLabel("Интервал (секунд):"))
        self.auto_save_interval = QSpinBox()
        self.auto_save_interval.setRange(10, 300)
        self.auto_save_interval.setValue(self.settings.config.get('auto_save', 30))
        auto_save_interval_layout.addWidget(self.auto_save_interval)
        auto_save_interval_layout.addStretch()

        auto_save_layout.addWidget(self.auto_save_check)
        auto_save_layout.addLayout(auto_save_interval_layout)
        auto_save_group.setLayout(auto_save_layout)
        layout.addWidget(auto_save_group)

        # Резервное копирование
        backup_group = QGroupBox("Резервное копирование")
        backup_layout = QVBoxLayout()

        self.backup_check = QCheckBox("Создавать резервные копии")
        self.backup_check.setChecked(self.settings.config.get('backup_enabled', True))

        backup_count_layout = QHBoxLayout()
        backup_count_layout.addWidget(QLabel("Хранить копий:"))
        self.backup_count = QSpinBox()
        self.backup_count.setRange(1, 50)
        self.backup_count.setValue(self.settings.config.get('backup_count', 10))
        backup_count_layout.addWidget(self.backup_count)
        backup_count_layout.addStretch()

        backup_layout.addWidget(self.backup_check)
        backup_layout.addLayout(backup_count_layout)
        backup_group.setLayout(backup_layout)
        layout.addWidget(backup_group)

        # Недавние проекты
        recent_group = QGroupBox("Недавние проекты")
        recent_layout = QHBoxLayout()
        recent_layout.addWidget(QLabel("Максимум проектов:"))
        self.max_recent = QSpinBox()
        self.max_recent.setRange(3, 30)
        self.max_recent.setValue(self.settings.config.get('max_recent', 10))
        recent_layout.addWidget(self.max_recent)
        recent_layout.addStretch()
        recent_group.setLayout(recent_layout)
        layout.addWidget(recent_group)

        layout.addStretch()
        widget.setLayout(layout)
        return widget

    def create_theme_tab(self) -> QWidget:
        """Создать вкладку настроек темы"""
        widget = QWidget()
        layout = QVBoxLayout()

        # Выбор темы
        theme_group = QGroupBox("Тема оформления")
        theme_layout = QVBoxLayout()

        # Список тем
        theme_selection_layout = QHBoxLayout()
        theme_selection_layout.addWidget(QLabel("Тема:"))

        self.theme_combo = QComboBox()
        for theme_id, theme_data in Theme.THEMES.items():
            self.theme_combo.addItem(theme_data['name'], theme_id)

        current_theme = self.settings.config.get('theme', 'dark')
        index = self.theme_combo.findData(current_theme)
        if index >= 0:
            self.theme_combo.setCurrentIndex(index)

        self.theme_combo.currentIndexChanged.connect(self.on_theme_changed)
        theme_selection_layout.addWidget(self.theme_combo)
        theme_selection_layout.addStretch()

        theme_layout.addLayout(theme_selection_layout)

        # Предпросмотр темы
        preview_label = QLabel("Предпросмотр:")
        theme_layout.addWidget(preview_label)

        preview_frame = QFrame()
        preview_frame.setFrameStyle(QFrame.Shape.Box)
        preview_frame.setMinimumHeight(100)

        preview_layout = QHBoxLayout()

        preview_btn = QPushButton("Кнопка")
        preview_layout.addWidget(preview_btn)

        preview_input = QLineEdit()
        preview_input.setText("Текст")
        preview_layout.addWidget(preview_input)

        preview_check = QCheckBox("Чекбокс")
        preview_layout.addWidget(preview_check)

        preview_layout.addStretch()
        preview_frame.setLayout(preview_layout)
        theme_layout.addWidget(preview_frame)

        # Обновляем стиль предпросмотра при смене темы
        self.preview_frame = preview_frame
        self.update_preview_style(current_theme)

        theme_group.setLayout(theme_layout)
        layout.addWidget(theme_group)

        # Пользовательские цвета
        colors_group = QGroupBox("Пользовательские цвета")
        colors_layout = QVBoxLayout()
        colors_layout.addWidget(QLabel("В разработке..."))
        colors_group.setLayout(colors_layout)
        layout.addWidget(colors_group)

        layout.addStretch()
        widget.setLayout(layout)
        return widget

    def create_editor_tab(self) -> QWidget:
        """Создать вкладку настроек редактора"""
        widget = QWidget()
        layout = QVBoxLayout()

        # Шрифт
        font_group = QGroupBox("Шрифт")
        font_layout = QGridLayout()

        font_layout.addWidget(QLabel("Семейство:"), 0, 0)
        self.font_combo = QFontComboBox()
        self.font_combo.setCurrentFont(QFont(self.settings.config.get('font_family', 'Consolas')))
        font_layout.addWidget(self.font_combo, 0, 1)

        font_layout.addWidget(QLabel("Размер:"), 1, 0)
        self.font_size = QSpinBox()
        self.font_size.setRange(8, 72)
        self.font_size.setValue(self.settings.config.get('font_size', 12))
        font_layout.addWidget(self.font_size, 1, 1)

        font_group.setLayout(font_layout)
        layout.addWidget(font_group)

        # Режимы
        mode_group = QGroupBox("Режимы")
        mode_layout = QVBoxLayout()

        self.vim_mode_check = QCheckBox("Vim-режим")
        self.vim_mode_check.setChecked(self.settings.config.get('vim_mode', True))
        mode_layout.addWidget(self.vim_mode_check)

        self.word_wrap_check = QCheckBox("Перенос слов")
        self.word_wrap_check.setChecked(self.settings.config.get('word_wrap', True))
        mode_layout.addWidget(self.word_wrap_check)

        mode_group.setLayout(mode_layout)
        layout.addWidget(mode_group)

        # Табуляция
        tab_group = QGroupBox("Табуляция")
        tab_layout = QHBoxLayout()
        tab_layout.addWidget(QLabel("Размер табуляции:"))
        self.tab_size = QSpinBox()
        self.tab_size.setRange(2, 8)
        self.tab_size.setValue(self.settings.config.get('tab_size', 4))
        tab_layout.addWidget(self.tab_size)
        tab_layout.addStretch()
        tab_group.setLayout(tab_layout)
        layout.addWidget(tab_group)

        layout.addStretch()
        widget.setLayout(layout)
        return widget

    def on_theme_changed(self, index: int):
        """При смене темы в комбобоксе"""
        theme_id = self.theme_combo.currentData()
        self.update_preview_style(theme_id)

    def update_preview_style(self, theme_id: str):
        """Обновить стиль предпросмотра"""
        theme = Theme(theme_id)
        self.preview_frame.setStyleSheet(f"""
            QFrame {{
                background-color: {theme.get('window_bg')};
                border: 1px solid {theme.get('border_color')};
            }}
            QPushButton {{
                background-color: {theme.get('button_bg')};
                color: {theme.get('button_fg')};
                border: 1px solid {theme.get('border_color')};
                padding: 5px;
            }}
            QLineEdit {{
                background-color: {theme.get('input_bg')};
                color: {theme.get('input_fg')};
                border: 1px solid {theme.get('input_border')};
                padding: 3px;
            }}
            QCheckBox {{
                color: {theme.get('window_fg')};
            }}
        """)

    def save_settings(self):
        """Сохранить настройки"""
        self.apply_settings()
        self.settings.save_config()
        self.accept()

    def apply_settings(self):
        """Применить настройки"""
        # Общие
        self.settings.config['auto_save_enabled'] = self.auto_save_check.isChecked()
        self.settings.config['auto_save'] = self.auto_save_interval.value()
        self.settings.config['backup_enabled'] = self.backup_check.isChecked()
        self.settings.config['backup_count'] = self.backup_count.value()
        self.settings.config['max_recent'] = self.max_recent.value()

        # Тема
        old_theme = self.settings.config.get('theme')
        new_theme = self.theme_combo.currentData()
        self.settings.config['theme'] = new_theme

        # Редактор
        self.settings.config['font_family'] = self.font_combo.currentFont().family()
        self.settings.config['font_size'] = self.font_size.value()
        self.settings.config['vim_mode'] = self.vim_mode_check.isChecked()
        self.settings.config['word_wrap'] = self.word_wrap_check.isChecked()
        self.settings.config['tab_size'] = self.tab_size.value()

        # Сигнал о смене темы
        if old_theme != new_theme:
            self.theme_changed.emit(new_theme)


class ChapterEditor(QWidget):
    """Редактор главы"""
    chapter_title_changed = pyqtSignal(str, str)  # chapter_id, new_title

    def __init__(self, project: Project, chapter: Chapter, main_window=None, parent=None):
        super().__init__(parent)
        self.main_window = main_window
        self.project = project
        self.chapter = chapter

        self.setWindowTitle(f"{chapter.title} - {project.name}")
        self.setMinimumSize(800, 600)

        layout = QVBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)

        # Верхняя панель
        toolbar = QToolBar()
        toolbar.setMovable(False)

        save_btn = QAction(QIcon.fromTheme('document-save'), 'Сохранить', self)
        save_btn.triggered.connect(self.save_current_chapter)
        toolbar.addAction(save_btn)

        toolbar.addSeparator()

        # Редактор названия
        toolbar.addWidget(QLabel("Название:"))
        self.title_editor = ChapterTitleEditor(chapter)
        self.title_editor.title_changed.connect(self.on_title_changed)
        toolbar.addWidget(self.title_editor)

        toolbar.addSeparator()

        # Статус
        self.status_combo = QComboBox()
        self.status_combo.addItems(['draft', 'review', 'done'])
        self.status_combo.setCurrentText(chapter.status)
        self.status_combo.currentTextChanged.connect(self.change_status)
        toolbar.addWidget(QLabel("Статус:"))
        toolbar.addWidget(self.status_combo)

        toolbar.addSeparator()

        # Статистика
        self.stats_label = QLabel()
        self.update_stats()
        toolbar.addWidget(self.stats_label)

        toolbar_widget = QWidget()
        toolbar_widget.setLayout(QHBoxLayout())
        toolbar_widget.layout().addWidget(toolbar)
        toolbar_widget.layout().addStretch()
        layout.addWidget(toolbar_widget)

        # Редактор
        self.editor = CodeEditor()
        self.editor.setPlainText(chapter.content)
        self.editor.textChanged.connect(self.on_text_changed)
        layout.addWidget(self.editor)

        # Vim статус
        layout.addWidget(self.editor.status_label)

        # Debounce для статистики
        self.stats_timer = QTimer()
        self.stats_timer.setSingleShot(True)
        self.stats_timer.timeout.connect(self._update_stats_debounced)
        self.stats_timer.setInterval(300)  # 300ms

        self.setLayout(layout)

    def on_title_changed(self, new_title: str):
        """Когда изменено название главы"""
        self.chapter.title = new_title
        self.setWindowTitle(f"{new_title} - {self.project.name}")
        self.chapter_title_changed.emit(self.chapter.id, new_title)
        self.save_current_chapter()

    def save_current_chapter(self):
        """Сохранить текущую главу"""
        try:
            self.chapter.content = self.editor.toPlainText()
            self.chapter.modified = datetime.now()
            if self.project.save():
                self.update_stats()
                if self.main_window and hasattr(self.main_window, 'statusBar'):
                    self.main_window.statusBar().showMessage("Сохранено", 2000)
        except Exception as e:
            print(f"Ошибка сохранения: {e}")

    def on_text_changed(self):
        """При изменении текста"""
        self.stats_timer.start()  # Запускаем debounce
        if hasattr(self.editor, 'document'):
            self.editor.document().setModified(True)

    def update_stats(self):
        """Обновить статистику (мгновенно)"""
        try:
            words = len(re.findall(r'\b\w+\b', self.editor.toPlainText()))
            chars = len(self.editor.toPlainText())
            self.stats_label.setText(f"Слов: {words} | Символов: {chars}")
        except Exception as e:
            print(f"Ошибка обновления статистики: {e}")

    def _update_stats_debounced(self):
        """Обновление статистики с debounce"""
        self.update_stats()

    def change_status(self, status):
        """Изменить статус"""
        try:
            self.chapter.status = status
            self.project.save()
        except Exception as e:
            print(f"Ошибка изменения статуса: {e}")


class MainWindow(QMainWindow):
    """Главное окно"""

    def __init__(self, ConfigClass):
        super().__init__()
        self.ConfigClass = ConfigClass
        self.current_project: Optional[Project] = None
        self.editors = {}
        self.settings = Settings()
        self.book_viewer = None
        self.temp_cover_path = None
        self.search_engine = None

        # Инициализация проверки обновлений
        self.updater = UpdateChecker("1.3")
        self.updater.update_found.connect(self._on_update_available)
        # Запускаем в фоне
        QTimer.singleShot(3000, self.updater.check_async)

        if ICON_PATH and os.path.exists(ICON_PATH):
            self.setWindowIcon(QIcon(ICON_PATH))
        icon_path = resource_path("icon.ico")
        if os.path.exists(icon_path):
            self.setWindowIcon(QIcon(icon_path))

        self.ConfigClass.init_dirs()

        self.setWindowTitle("Читалка-писалка Pro")
        self.setGeometry(100, 100, 1400, 900)

        self.mdi_area = QMdiArea()
        self.setCentralWidget(self.mdi_area)

        self.create_menu()
        self.create_toolbar()
        self.create_dock_widgets()

        self.statusBar().showMessage("Готов")

        # Применяем сохраненную тему
        self.apply_theme(self.settings.config.get('theme', 'dark'))

        QTimer.singleShot(100, self.show_start_page)

    def create_menu(self):
        """Создать меню"""
        menubar = self.menuBar()
        menubar.clear()

        # Файл
        file_menu = menubar.addMenu("&Файл")

        new_project = QAction(QIcon.fromTheme('document-new'), "&Новый проект", self)
        new_project.setShortcut("Ctrl+N")
        new_project.triggered.connect(self.new_project)
        file_menu.addAction(new_project)

        open_project_action = QAction(QIcon.fromTheme('document-open'), "&Открыть проект...", self)
        open_project_action.setShortcut("Ctrl+O")
        open_project_action.triggered.connect(self.open_project_dialog)
        file_menu.addAction(open_project_action)

        # Предпросмотр
        open_viewer_action = QAction(QIcon.fromTheme('document-open'), "Предпросмотр книги...", self)
        open_viewer_action.triggered.connect(self.show_book_viewer)
        file_menu.addAction(open_viewer_action)

        self.recent_menu = file_menu.addMenu("&Недавние проекты")
        self.update_recent_menu()

        file_menu.addSeparator()

        # НОВОЕ: Импорт файлов
        import_menu = file_menu.addMenu("📥 Импорт")

        import_txt = QAction("TXT", self)
        import_txt.triggered.connect(lambda: self.import_file('txt'))
        import_menu.addAction(import_txt)

        import_fb2 = QAction("FB2", self)
        import_fb2.triggered.connect(lambda: self.import_file('fb2'))
        import_menu.addAction(import_fb2)

        import_docx = QAction("DOCX", self)
        import_docx.triggered.connect(lambda: self.import_file('docx'))
        import_menu.addAction(import_docx)

        import_pdf = QAction("PDF (текст)", self)
        import_pdf.triggered.connect(lambda: self.import_file('pdf'))
        import_menu.addAction(import_pdf)

        file_menu.addAction(import_menu.menuAction())

        file_menu.addSeparator()

        save_all = QAction(QIcon.fromTheme('document-save-all'), "&Сохранить всё", self)
        save_all.setShortcut("Ctrl+Shift+S")
        save_all.triggered.connect(self.save_all)
        file_menu.addAction(save_all)

        file_menu.addSeparator()

        # Подменю для обложки
        cover_menu = file_menu.addMenu("📷 &Обложка")

        add_cover_action = QAction("Добавить/Изменить обложку", self)
        add_cover_action.triggered.connect(self.manage_cover)
        cover_menu.addAction(add_cover_action)

        remove_cover_action = QAction("Удалить обложку", self)
        remove_cover_action.triggered.connect(self.remove_cover)
        cover_menu.addAction(remove_cover_action)

        cover_menu.addSeparator()

        view_cover_action = QAction("Просмотреть обложку", self)
        view_cover_action.triggered.connect(self.view_cover)
        cover_menu.addAction(view_cover_action)

        file_menu.addSeparator()

        export_menu = file_menu.addMenu("&Экспорт")

        export_fb2 = QAction("FB2", self)
        export_fb2.triggered.connect(lambda: self.export_project('fb2'))
        export_menu.addAction(export_fb2)

        export_pdf = QAction("PDF", self)
        export_pdf.triggered.connect(lambda: self.export_project('pdf'))
        export_menu.addAction(export_pdf)

        export_txt = QAction("TXT", self)
        export_txt.triggered.connect(lambda: self.export_project('txt'))
        export_menu.addAction(export_txt)

        export_docx = QAction("DOCX", self)
        export_docx.triggered.connect(lambda: self.export_project('docx'))
        export_menu.addAction(export_docx)

        file_menu.addSeparator()

        exit_action = QAction("&Выход", self)
        exit_action.setShortcut("Ctrl+Q")
        exit_action.triggered.connect(self.close)
        file_menu.addAction(exit_action)

        # Настройки
        settings_menu = menubar.addMenu("&Настройки")

        theme_menu = settings_menu.addMenu("&Тема оформления")

        for theme_id, theme_data in Theme.THEMES.items():
            action = QAction(theme_data['name'], self)
            action.setData(theme_id)
            action.triggered.connect(lambda checked, t=theme_id: self.change_theme(t))
            if theme_id == self.settings.config.get('theme'):
                action.setCheckable(True)
                action.setChecked(True)
            theme_menu.addAction(action)

        theme_menu.addSeparator()

        settings_action = QAction("&Настройки...", self)
        settings_action.setShortcut("Ctrl+,")
        settings_action.triggered.connect(self.show_settings)
        settings_menu.addAction(settings_action)

        # Справка
        help_menu = menubar.addMenu("&Справка")

        check_update_action = QAction("Проверить обновления", self)
        check_update_action.triggered.connect(self._manual_check)
        help_menu.addAction(check_update_action)

        help_menu.addSeparator()

        about_action = QAction("&О программе", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)

    # ===== ИМПОРТ ФАЙЛОВ =====

    def import_file(self, format: str):
        """Импорт файла в проект"""
        if not self.current_project:
            QMessageBox.warning(self, "Предупреждение", "Сначала откройте проект")
            return

        file_path, _ = QFileDialog.getOpenFileName(
            self,
            f"Импорт {format.upper()}",
            os.path.expanduser("~"),
            f"*.{format}"
        )

        if not file_path:
            return

        try:
            # Создаем новую главу
            chapter = self.current_project.add_chapter(f"Импорт: {os.path.basename(file_path)}")
            content = ""

            if format == 'txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                chapter.content = content

            elif format == 'fb2':
                content = self._parse_fb2_to_text(file_path)
                chapter.content = content

            elif format == 'docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    content = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    chapter.content = content
                except ImportError:
                    QMessageBox.warning(self, "Ошибка", "Установите python-docx для импорта DOCX")
                    return

            elif format == 'pdf':
                if PYPDF2_SUPPORT:
                    try:
                        with open(file_path, 'rb') as f:
                            reader = PyPDF2.PdfReader(f)
                            text_parts = []
                            for page in reader.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    text_parts.append(page_text)
                            content = "\n\n".join(text_parts)
                            chapter.content = content
                    except Exception as e:
                        QMessageBox.warning(self, "Ошибка", f"Не удалось извлечь текст из PDF: {e}")
                        return
                else:
                    QMessageBox.warning(self, "Ошибка", "Установите PyPDF2 для импорта PDF")
                    return

            self.current_project.save()

            # Добавляем в дерево
            item = ChapterTreeItem(chapter)
            self.project_tree.addTopLevelItem(item)

            QMessageBox.information(
                self,
                "Импорт завершен",
                f"Файл импортирован как глава:\n{chapter.title}\n\nСлов: {chapter.word_count()}"
            )

        except Exception as e:
            QMessageBox.critical(self, "Ошибка импорта", str(e))

    def _parse_fb2_to_text(self, file_path: str) -> str:
        """Парсинг FB2 в простой текст"""
        try:
            it = ET.iterparse(file_path)
            for _, el in it:
                if '}' in el.tag:
                    el.tag = el.tag.split('}', 1)[1]
            root = it.root

            text_parts = []
            body = root.find("body")
            if body is not None:
                for section in body.findall("section"):
                    title = section.find("title")
                    if title is not None:
                        p_title = title.find("p")
                        if p_title is not None and p_title.text:
                            text_parts.append(f"\n\n=== {p_title.text} ===\n")

                    for p in section.findall("p"):
                        if p.text:
                            text_parts.append(p.text)

            return "\n\n".join(text_parts)
        except Exception as e:
            return f"Ошибка парсинга FB2: {e}"

    # ===== ВСПОМОГАТЕЛЬНЫЕ МЕТОДЫ =====

    def _on_update_available(self, version: str, url: str):
        """Обработчик найденного обновления"""
        QMetaObject.invokeMethod(
            self,
            "_show_update_notification",
            Qt.ConnectionType.QueuedConnection,
            Q_ARG(str, version),
            Q_ARG(str, url)
        )

    @pyqtSlot(str, str)
    def _show_update_notification(self, version: str, url: str):
        """Показать уведомление об обновлении (в основном потоке)"""
        reply = QMessageBox.question(
            self,
            "Доступно обновление!",
            f"Версия {version} уже доступна.\nХотите перейти на страницу загрузки?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            QDesktopServices.openUrl(QUrl(url))

    def _manual_check(self):
        """Ручная проверка обновлений"""
        self.statusBar().showMessage("Проверка обновлений...", 2000)
        self.updater.check_async()
        QTimer.singleShot(3000, lambda: self.statusBar().showMessage("Проверка завершена", 2000))

    def show_book_viewer(self):
        """Показать окно предпросмотра"""
        if self.book_viewer is None:
            self.book_viewer = BookViewer(self)
        self.book_viewer.show()
        self.book_viewer.raise_()
        self.book_viewer.activateWindow()

    def manage_cover(self):
        """Управление обложкой"""
        if not self.current_project:
            QMessageBox.warning(self, "Предупреждение", "Сначала откройте проект")
            return

        dialog = CoverDialog(self.current_project, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.update_stats_display()
            self.statusBar().showMessage("Обложка обновлена", 2000)

    def remove_cover(self):
        """Удалить обложку"""
        if not self.current_project:
            return

        if not self.current_project.cover_path:
            QMessageBox.information(self, "Информация", "У проекта нет обложки")
            return

        reply = QMessageBox.question(
            self,
            "Подтверждение",
            "Удалить обложку?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.current_project.remove_cover()
            self.current_project.save()
            self.update_stats_display()
            self.statusBar().showMessage("Обложка удалена", 2000)

    def view_cover(self):
        """Просмотреть обложку в отдельном окне"""
        if not self.current_project or not self.current_project.cover_path:
            QMessageBox.information(self, "Информация", "У проекта нет обложки")
            return

        pixmap = self.current_project.get_cover_pixmap(800, 600)
        if pixmap:
            dialog = QDialog(self)
            dialog.setWindowTitle("Обложка книги")
            layout = QVBoxLayout()

            label = QLabel()
            label.setPixmap(pixmap)
            label.setAlignment(Qt.AlignmentFlag.AlignCenter)

            scroll = QScrollArea()
            scroll.setWidget(label)
            scroll.setWidgetResizable(True)

            layout.addWidget(scroll)

            close_btn = QPushButton("Закрыть")
            close_btn.clicked.connect(dialog.accept)
            layout.addWidget(close_btn)

            dialog.setLayout(layout)
            dialog.resize(900, 700)
            dialog.exec()

    def change_theme(self, theme_name: str):
        """Сменить тему"""
        self.settings.update_theme(theme_name)
        self.apply_theme(theme_name)
        self.create_menu()

    def apply_theme(self, theme_name: str):
        """Применить тему"""
        theme = Theme(theme_name)
        self.setStyleSheet(theme.get_stylesheet())

        for sub in self.mdi_area.subWindowList():
            if isinstance(sub.widget(), ChapterEditor):
                sub.widget().editor.setStyleSheet("")
                sub.widget().editor.setStyleSheet(theme.get_stylesheet())

    def show_settings(self):
        """Показать диалог настроек"""
        dialog = SettingsDialog(self.settings, self)
        dialog.theme_changed.connect(self.change_theme)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            self.apply_theme(self.settings.config.get('theme'))
            self.statusBar().showMessage("Настройки сохранены", 2000)

    def show_about(self):
        """Показать информацию о программе"""
        QMessageBox.about(
            self,
            "О программе",
            """<h1>📚 Читалка-писалка Pro</h1>
            <p>Версия: 1.2</p>
            <p>Профессиональный инструмент для писателей</p>
            <br>
            <p><b>Возможности:</b></p>
            <ul>
                <li>Многодокументный интерфейс</li>
                <li>Дерево проекта с Drag&Drop</li>
                <li>Система тем оформления</li>
                <li>Поддержка обложек книг</li>
                <li>Статистика в реальном времени</li>
                <li>Экспорт в FB2/PDF/TXT/DOCX с обложкой</li>
                <li>Импорт из FB2/PDF/TXT/DOCX</li>
                <li>Просмотр книг как в FBReader</li>
                <li>Vim-режим</li>
                <li>Система недавних проектов</li>
                <li>Автосохранение и бэкапы</li>
                <li>Автоматическая проверка обновлений</li>
                <li>Улучшенный поиск по проекту</li>
            </ul>
            <br>
            <p>© 2026 Читалка-писалка Pro</p>"""
        )

    # ===== РАБОТА С НЕДАВНИМИ ПРОЕКТАМИ =====

    def update_recent_menu(self):
        """Обновить меню недавних проектов"""
        self.recent_menu.clear()

        recent = self.settings.get_recent_projects()
        if not recent:
            action = QAction("Нет недавних проектов", self)
            action.setEnabled(False)
            self.recent_menu.addAction(action)
            return

        for project in recent:
            name = project['name']
            path = project['path']
            action = QAction(name, self)
            action.setData(path)
            action.triggered.connect(lambda checked, p=path: self.open_project_path(p))
            self.recent_menu.addAction(action)

        self.recent_menu.addSeparator()
        clear_action = QAction("Очистить список", self)
        clear_action.triggered.connect(self.clear_recent)
        self.recent_menu.addAction(clear_action)

    def clear_recent(self):
        """Очистить список недавних проектов"""
        self.settings.recent_projects = []
        self.settings.save_recent()
        self.update_recent_menu()

    # ===== ТУЛБАР =====

    def create_toolbar(self):
        """Создать тулбар"""
        toolbar = self.addToolBar("Основная")

        new_btn = QAction(QIcon.fromTheme('document-new'), "Новый проект", self)
        new_btn.triggered.connect(self.new_project)
        toolbar.addAction(new_btn)

        open_btn = QAction(QIcon.fromTheme('document-open'), "Открыть", self)
        open_btn.triggered.connect(self.open_project_dialog)
        toolbar.addAction(open_btn)

        toolbar.addSeparator()

        save_btn = QAction(QIcon.fromTheme('document-save'), "Сохранить всё", self)
        save_btn.triggered.connect(self.save_all)
        toolbar.addAction(save_btn)

        toolbar.addSeparator()

        add_chapter_btn = QAction(QIcon.fromTheme('list-add'), "Добавить главу", self)
        add_chapter_btn.triggered.connect(self.add_chapter)
        toolbar.addAction(add_chapter_btn)

        toolbar.addSeparator()

        cover_btn = QAction(QIcon.fromTheme('image-x-generic'), "Обложка", self)
        cover_btn.triggered.connect(self.manage_cover)
        toolbar.addAction(cover_btn)

        toolbar.addSeparator()

        settings_btn = QAction(QIcon.fromTheme('preferences-system'), "Настройки", self)
        settings_btn.triggered.connect(self.show_settings)
        toolbar.addAction(settings_btn)

    # ===== DOCK WIDGETS =====

    def create_dock_widgets(self):
        """Создать закрепляемые панели"""
        # Дерево проекта
        self.project_dock = QDockWidget("Проект", self)
        self.project_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea |
                                          Qt.DockWidgetArea.RightDockWidgetArea)

        self.project_tree = QTreeWidget()
        self.project_tree.setHeaderLabel("Главы")
        self.project_tree.setDragEnabled(True)
        self.project_tree.setAcceptDrops(True)
        self.project_tree.setDragDropMode(QAbstractItemView.DragDropMode.InternalMove)
        self.project_tree.itemDoubleClicked.connect(self.open_chapter)
        self.project_tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.project_tree.customContextMenuRequested.connect(self.show_chapter_context_menu)

        self.project_dock.setWidget(self.project_tree)
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.project_dock)

        # Панель статистики
        self.stats_dock = QDockWidget("Статистика", self)
        self.stats_widget = QWidget()
        self.stats_layout = QVBoxLayout()

        # Отображение обложки в статистике
        self.cover_preview = QLabel()
        self.cover_preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.cover_preview.setMinimumHeight(150)
        self.cover_preview.setStyleSheet("border: 1px solid gray; margin: 5px;")
        self.stats_layout.addWidget(self.cover_preview)

        self.stats_text = QLabel()
        self.stats_text.setWordWrap(True)
        self.stats_text.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.stats_text.setOpenExternalLinks(False)

        self.stats_layout.addWidget(self.stats_text)
        self.stats_layout.addStretch()

        self.stats_widget.setLayout(self.stats_layout)
        self.stats_dock.setWidget(self.stats_widget)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.stats_dock)

        # Панель поиска
        self.search_dock = QDockWidget("Поиск", self)
        search_widget = QWidget()
        search_layout = QVBoxLayout()

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Поиск по проекту...")
        self.search_input.textChanged.connect(self.search_text)

        self.search_results = QListWidget()
        self.search_results.itemDoubleClicked.connect(self.goto_search_result)

        search_layout.addWidget(self.search_input)
        search_layout.addWidget(self.search_results)

        search_widget.setLayout(search_layout)
        self.search_dock.setWidget(search_widget)
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, self.search_dock)

        self.search_dock.hide()

    # ===== КОНТЕКСТНОЕ МЕНЮ ГЛАВЫ =====

    def show_chapter_context_menu(self, position):
        """Показать контекстное меню для главы"""
        item = self.project_tree.itemAt(position)
        if not isinstance(item, ChapterTreeItem):
            return

        menu = QMenu()

        rename_action = menu.addAction("Переименовать")
        delete_action = menu.addAction("Удалить")
        menu.addSeparator()

        status_menu = menu.addMenu("Изменить статус")
        draft_action = status_menu.addAction("Черновик")
        review_action = status_menu.addAction("На ревью")
        done_action = status_menu.addAction("Готово")

        action = menu.exec(self.project_tree.viewport().mapToGlobal(position))

        if action == rename_action:
            self.rename_chapter(item)
        elif action == delete_action:
            self.delete_chapter(item)
        elif action == draft_action:
            self.change_chapter_status(item, 'draft')
        elif action == review_action:
            self.change_chapter_status(item, 'review')
        elif action == done_action:
            self.change_chapter_status(item, 'done')

    def rename_chapter(self, item: ChapterTreeItem):
        """Переименовать главу"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Переименовать главу")

        layout = QVBoxLayout()
        layout.addWidget(QLabel("Новое название:"))

        title_input = QLineEdit(item.chapter.title)
        layout.addWidget(title_input)

        btn_layout = QHBoxLayout()
        ok_btn = QPushButton("ОК")
        cancel_btn = QPushButton("Отмена")

        btn_layout.addWidget(ok_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

        dialog.setLayout(layout)

        def rename():
            new_title = title_input.text().strip()
            if new_title and new_title != item.chapter.title:
                if self.current_project.update_chapter_title(item.chapter.id, new_title):
                    self.current_project.save()
                    item.chapter.title = new_title
                    item.update_display()

                    if item.chapter.id in self.editors:
                        editor = self.editors[item.chapter.id].widget()
                        if isinstance(editor, ChapterEditor):
                            editor.setWindowTitle(f"{new_title} - {self.current_project.name}")
                            editor.title_editor.title_edit.setText(new_title)

                    self.statusBar().showMessage("Глава переименована", 2000)
            dialog.accept()

        ok_btn.clicked.connect(rename)
        cancel_btn.clicked.connect(dialog.reject)

        dialog.exec()

    def delete_chapter(self, item: ChapterTreeItem):
        """Удалить главу"""
        reply = QMessageBox.question(
            self,
            "Подтверждение",
            f"Удалить главу '{item.chapter.title}'?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            if item.chapter.id in self.editors:
                self.editors[item.chapter.id].close()

            self.current_project.delete_chapter(item.chapter.id)
            self.current_project.save()

            index = self.project_tree.indexOfTopLevelItem(item)
            self.project_tree.takeTopLevelItem(index)

            for i in range(self.project_tree.topLevelItemCount()):
                tree_item = self.project_tree.topLevelItem(i)
                if isinstance(tree_item, ChapterTreeItem):
                    tree_item.chapter.order = i
                    tree_item.update_display()

            self.update_stats_display()
            self.statusBar().showMessage("Глава удалена", 2000)

    def change_chapter_status(self, item: ChapterTreeItem, status: str):
        """Изменить статус главы"""
        item.chapter.status = status
        item.update_display()
        self.current_project.save()

        if item.chapter.id in self.editors:
            editor = self.editors[item.chapter.id].widget()
            if isinstance(editor, ChapterEditor):
                editor.status_combo.setCurrentText(status)

        self.update_stats_display()
        self.statusBar().showMessage(f"Статус изменен на {status}", 2000)

    # ===== СТАРТОВАЯ СТРАНИЦА =====

    def show_start_page(self):
        """Показать стартовую страницу"""
        try:
            self.mdi_area.closeAllSubWindows()
            self.project_tree.clear()
            self.current_project = None

            start_widget = QWidget()
            layout = QVBoxLayout()
            layout.setSpacing(20)

            title = QLabel("📚 Читалка-писалка Pro")
            title.setStyleSheet("font-size: 32px; font-weight: bold; margin: 30px;")
            title.setAlignment(Qt.AlignmentFlag.AlignCenter)
            layout.addWidget(title)

            btn_layout = QHBoxLayout()
            btn_layout.setSpacing(20)

            new_btn = QPushButton("➕ Создать новый проект")
            new_btn.setMinimumHeight(60)
            new_btn.setMinimumWidth(200)
            new_btn.clicked.connect(self.new_project)
            btn_layout.addWidget(new_btn)

            open_btn = QPushButton("📂 Открыть существующий")
            open_btn.setMinimumHeight(60)
            open_btn.setMinimumWidth(200)
            open_btn.clicked.connect(self.open_project_dialog)
            btn_layout.addWidget(open_btn)

            layout.addLayout(btn_layout)

            recent = self.settings.get_recent_projects()
            if recent:
                recent_label = QLabel("Недавние проекты:")
                recent_label.setStyleSheet("font-size: 18px; font-weight: bold; margin-top: 30px;")
                layout.addWidget(recent_label)

                for project in recent[:5]:
                    project_widget = self.create_recent_project_widget(project)
                    layout.addWidget(project_widget)

            layout.addStretch()
            start_widget.setLayout(layout)

            scroll = QScrollArea()
            scroll.setWidget(start_widget)
            scroll.setWidgetResizable(True)
            scroll.setStyleSheet("QScrollArea { border: none; }")

            sub = self.mdi_area.addSubWindow(scroll)
            sub.setWindowTitle("Стартовая страница")
            sub.show()

        except Exception as e:
            print(f"Ошибка при показе стартовой страницы: {e}")

    def create_recent_project_widget(self, project: Dict) -> QWidget:
        """Создать виджет недавнего проекта"""
        widget = QWidget()
        layout = QHBoxLayout()
        layout.setContentsMargins(10, 5, 10, 5)

        icon_label = QLabel("📁")
        icon_label.setStyleSheet("font-size: 24px;")
        layout.addWidget(icon_label)

        info_layout = QVBoxLayout()
        name_label = QLabel(project['name'])
        name_label.setStyleSheet("font-weight: bold;")

        path_label = QLabel(project['path'])
        path_label.setStyleSheet("color: gray; font-size: 10px;")

        info_layout.addWidget(name_label)
        info_layout.addWidget(path_label)
        layout.addLayout(info_layout, 1)

        open_btn = QPushButton("Открыть")
        open_btn.clicked.connect(lambda: self.open_project_path(project['path']))
        layout.addWidget(open_btn)

        widget.setLayout(layout)
        return widget

    # ===== СОЗДАНИЕ/ОТКРЫТИЕ ПРОЕКТА =====

    def new_project(self):
        """Создать новый проект"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Новый проект")
        dialog.setMinimumWidth(500)

        layout = QVBoxLayout()
        layout.setSpacing(15)

        layout.addWidget(QLabel("Название проекта:"))
        name_input = QLineEdit()
        name_input.setPlaceholderText("Введите название проекта")
        layout.addWidget(name_input)

        layout.addWidget(QLabel("Автор:"))
        author_input = QLineEdit()
        author_input.setPlaceholderText("Ваше имя")
        layout.addWidget(author_input)

        layout.addWidget(QLabel("Папка для сохранения:"))
        path_layout = QHBoxLayout()

        path_input = QLineEdit()
        default_path = os.path.join(self.ConfigClass.PROJECTS_DIR, "Новый проект")
        path_input.setText(default_path)
        path_input.setReadOnly(True)

        browse_btn = QPushButton("Обзор...")
        browse_btn.clicked.connect(lambda: self.browse_project_folder(path_input))

        path_layout.addWidget(path_input)
        path_layout.addWidget(browse_btn)
        layout.addLayout(path_layout)

        # Опция добавить обложку сразу
        cover_layout = QHBoxLayout()
        cover_layout.addWidget(QLabel("Обложка (можно добавить позже):"))
        self.temp_cover_path = None
        self.temp_cover_label = QLabel("Не выбрана")
        cover_layout.addWidget(self.temp_cover_label)
        select_cover_btn = QPushButton("Выбрать...")
        select_cover_btn.clicked.connect(self.select_temp_cover)
        cover_layout.addWidget(select_cover_btn)
        layout.addLayout(cover_layout)

        btn_layout = QHBoxLayout()
        create_btn = QPushButton("Создать")
        create_btn.setDefault(True)
        cancel_btn = QPushButton("Отмена")

        btn_layout.addWidget(create_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

        dialog.setLayout(layout)

        def create():
            name = name_input.text().strip()
            if not name:
                QMessageBox.warning(dialog, "Ошибка", "Введите название проекта")
                return

            safe_name = self.ConfigClass.sanitize_folder_name(name)

            path = path_input.text().strip()
            if not path or path == default_path:
                path = os.path.join(self.ConfigClass.PROJECTS_DIR, safe_name)

            if os.path.exists(path):
                reply = QMessageBox.question(
                    dialog,
                    "Подтверждение",
                    f"Папка уже существует. Открыть существующий проект?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
                )
                if reply == QMessageBox.StandardButton.Yes:
                    dialog.accept()
                    self.open_project_path(path)
                return

            project = Project(name, path)
            if author_input.text():
                project.settings['author'] = author_input.text()

            # Добавляем обложку если выбрана
            if self.temp_cover_path:
                project.set_cover(self.temp_cover_path)

            if project.save():
                dialog.accept()
                self.settings.add_recent_project(path, name)
                self.update_recent_menu()
                self.open_project_path(path)
            else:
                QMessageBox.critical(dialog, "Ошибка",
                                     "Не удалось создать проект. Проверьте права доступа.")

        create_btn.clicked.connect(create)
        cancel_btn.clicked.connect(dialog.reject)

        dialog.exec()

    def select_temp_cover(self):
        """Выбрать временную обложку при создании проекта"""
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Выберите изображение для обложки",
            os.path.expanduser("~"),
            "Изображения (*.png *.jpg *.jpeg *.bmp *.gif)"
        )
        if path:
            self.temp_cover_path = path
            self.temp_cover_label.setText(os.path.basename(path))

    def browse_project_folder(self, line_edit: QLineEdit):
        """Выбрать папку для проекта"""
        path = QFileDialog.getExistingDirectory(
            self,
            "Выберите папку для проекта",
            self.ConfigClass.PROJECTS_DIR
        )
        if path:
            line_edit.setText(path)

    def open_project_dialog(self):
        """Диалог открытия проекта"""
        path = QFileDialog.getExistingDirectory(
            self,
            "Выберите папку проекта",
            self.ConfigClass.PROJECTS_DIR
        )

        if path:
            self.open_project_path(path)

    def open_project_path(self, path: str):
        """Открыть проект по пути"""
        try:
            project = Project.load(path)
            if project:
                self.load_project(project)
                self.settings.add_recent_project(path, project.name)
                self.update_recent_menu()
            else:
                QMessageBox.critical(self, "Ошибка",
                                     "Не удалось загрузить проект. Файл проекта поврежден или отсутствует.")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть проект: {e}")

    def load_project(self, project: Project):
        """Загрузить проект в интерфейс"""
        try:
            self.current_project = project
            self.setWindowTitle(f"Читалка-писалка Pro - {project.name}")

            # Инициализируем поиск
            self.search_engine = SearchEngine(project)
            self.search_engine.result_found.connect(self._add_search_result)

            self.mdi_area.closeAllSubWindows()

            self.project_tree.clear()
            for chapter in project.chapters:
                item = ChapterTreeItem(chapter)
                self.project_tree.addTopLevelItem(item)

            self.update_stats_display()

            self.project_dock.show()
            self.stats_dock.show()

            self.statusBar().showMessage(f"Проект '{project.name}' загружен")
        except Exception as e:
            print(f"Ошибка загрузки проекта в интерфейс: {e}")

    # ===== ПОИСК =====

    def _add_search_result(self, result: dict):
        """Добавить результат поиска в список"""
        item = QListWidgetItem(f"{result['chapter_title']}: ...{result['context']}...")
        item.setData(Qt.ItemDataRole.UserRole, result)
        self.search_results.addItem(item)

    def search_text(self, text: str):
        """Поиск по проекту"""
        if not self.current_project:
            return

        self.search_results.clear()
        if len(text) < 2:
            self.search_dock.hide()
            return

        self.search_dock.show()
        if self.search_engine:
            self.search_engine.search(text)

    def goto_search_result(self, item: QListWidgetItem):
        """Перейти к результату поиска"""
        try:
            result = item.data(Qt.ItemDataRole.UserRole)
            if not result:
                return

            chapter_id = result['chapter_id']
            position = result['position']

            for sub in self.mdi_area.subWindowList():
                if isinstance(sub.widget(), ChapterEditor):
                    if sub.widget().chapter.id == chapter_id:
                        sub.show()
                        sub.raise_()
                        editor = sub.widget().editor
                        cursor = editor.textCursor()
                        cursor.setPosition(position)
                        cursor.setPosition(
                            position + len(self.search_input.text()),
                            QTextCursor.MoveMode.KeepAnchor
                        )
                        editor.setTextCursor(cursor)
                        editor.setFocus()
                        return

            chapter = self.current_project.get_chapter(chapter_id)
            if chapter:
                temp_item = ChapterTreeItem(chapter)
                self.open_chapter(temp_item)

        except Exception as e:
            print(f"Ошибка перехода к результату поиска: {e}")

    # ===== СТАТИСТИКА =====

    def update_stats_display(self):
        """Обновить отображение статистики"""
        if not self.current_project:
            return

        try:
            stats = self.current_project.get_stats()

            # Обновляем предпросмотр обложки
            pixmap = self.current_project.get_cover_pixmap(180, 250)
            if pixmap:
                self.cover_preview.setPixmap(pixmap)
                self.cover_preview.setToolTip("Обложка книги")
            else:
                self.cover_preview.setText("📖\nНет обложки")
                self.cover_preview.setStyleSheet("border: 1px solid gray; margin: 5px; font-size: 14px;")

            text = f"""
            <b>Статистика проекта</b><br><br>
            📚 Глав: {stats['chapters']}<br>
            📝 Слов: {stats['words']:,}<br>
            🔤 Символов: {stats['chars']:,}<br>
            <br>
            <b>Статусы:</b><br>
            📝 Черновиков: {stats['drafts']}<br>
            🔍 На ревью: {stats['review']}<br>
            ✅ Готово: {stats['done']}<br>
            <br>
            <b>Обложка:</b><br>
            {'✅ Есть' if stats['has_cover'] else '❌ Нет'}<br>
            <br>
            <b>Даты:</b><br>
            ✨ Создан: {stats['created']}<br>
            📅 Изменён: {stats['modified']}<br>
            """

            self.stats_text.setText(text)
        except Exception as e:
            print(f"Ошибка обновления статистики: {e}")

    # ===== ГЛАВЫ =====

    def add_chapter(self):
        """Добавить главу"""
        if not self.current_project:
            QMessageBox.warning(self, "Предупреждение", "Сначала откройте проект")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("Новая глава")

        layout = QVBoxLayout()

        layout.addWidget(QLabel("Название главы:"))
        title_input = QLineEdit(f"Глава {len(self.current_project.chapters) + 1}")
        layout.addWidget(title_input)

        btn_layout = QHBoxLayout()
        create_btn = QPushButton("Создать")
        cancel_btn = QPushButton("Отмена")

        btn_layout.addWidget(create_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

        dialog.setLayout(layout)

        def create():
            title = title_input.text().strip()
            chapter = self.current_project.add_chapter(title)
            self.current_project.save()

            item = ChapterTreeItem(chapter)
            self.project_tree.addTopLevelItem(item)

            dialog.accept()
            self.open_chapter(item)

        create_btn.clicked.connect(create)
        cancel_btn.clicked.connect(dialog.reject)

        dialog.exec()

    def open_chapter(self, item):
        """Открыть главу"""
        if not isinstance(item, ChapterTreeItem):
            item = self.project_tree.currentItem()
            if not isinstance(item, ChapterTreeItem):
                return

        chapter = item.chapter

        if chapter.id in self.editors:
            self.mdi_area.setActiveSubWindow(self.editors[chapter.id])
            return

        editor = ChapterEditor(self.current_project, chapter, self)
        editor.chapter_title_changed.connect(self.on_chapter_title_changed)

        sub = self.mdi_area.addSubWindow(editor)
        sub.setWindowTitle(chapter.title)
        sub.show()

        self.editors[chapter.id] = sub
        sub.destroyed.connect(lambda: self.editors.pop(chapter.id, None))

    def on_chapter_title_changed(self, chapter_id: str, new_title: str):
        """Когда изменено название главы в редакторе"""
        for i in range(self.project_tree.topLevelItemCount()):
            item = self.project_tree.topLevelItem(i)
            if isinstance(item, ChapterTreeItem) and item.chapter.id == chapter_id:
                item.update_display()
                break

    def save_all(self):
        """Сохранить всё"""
        if not self.current_project:
            return

        try:
            for sub in self.mdi_area.subWindowList():
                if isinstance(sub.widget(), ChapterEditor):
                    sub.widget().save_current_chapter()

            if self.current_project.save():
                self.update_stats_display()
                self.statusBar().showMessage("Всё сохранено", 2000)
        except Exception as e:
            print(f"Ошибка сохранения: {e}")

    # ===== ЭКСПОРТ =====

    def export_project(self, format: str):
        """Экспорт проекта"""
        if not self.current_project:
            QMessageBox.warning(self, "Предупреждение", "Сначала откройте проект")
            return

        default_dir = os.path.dirname(self.current_project.path)
        ext = format
        default_name = f"{self.current_project.name}.{ext}"

        path, _ = QFileDialog.getSaveFileName(
            self,
            f"Экспорт в {format.upper()}",
            os.path.join(default_dir, default_name),
            f"*.{ext}"
        )

        if not path:
            return

        if not path.lower().endswith(f'.{ext}'):
            path += f'.{ext}'

        try:
            if format == 'fb2':
                if FB2Exporter:
                    FB2Exporter.export(self.current_project, path)
                else:
                    self.export_fb2_fallback(path)
            elif format == 'docx':
                if DOCXExporter:
                    DOCXExporter.export(self.current_project, path)
                else:
                    QMessageBox.warning(self, "Ошибка", "Модуль DOCX не установлен. Установите python-docx")
                    return
            elif format == 'txt':
                if TXTExporter:
                    TXTExporter.export(self.current_project, path)
                else:
                    self.export_txt_fallback(path)
            elif format == 'pdf':
                if PDFExporter:
                    PDFExporter.export(self.current_project, path)
                else:
                    self.export_pdf_fallback(path)

            reply = QMessageBox.question(
                self,
                "Экспорт завершен",
                f"Проект экспортирован в:\n{path}\n\nОткрыть файл?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )

            if reply == QMessageBox.StandardButton.Yes:
                QDesktopServices.openUrl(QUrl.fromLocalFile(path))

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка экспорта: {e}")

    # === FALLBACK методы экспорта (если модули не загружены) ===

    def export_fb2_fallback(self, path: str):
        """Экспорт в FB2 (встроенный метод)"""
        try:
            root = ET.Element("FictionBook")
            root.set("xmlns", "http://www.gribuser.ru/xml/fictionbook/2.0")
            root.set("xmlns:l", "http://www.w3.org/1999/xlink")

            description = ET.SubElement(root, "description")
            title_info = ET.SubElement(description, "title-info")

            if self.current_project.cover_data:
                coverpage = ET.SubElement(title_info, "coverpage")
                image = ET.SubElement(coverpage, "image")
                image.set("{http://www.w3.org/1999/xlink}href", "#cover.jpg")
                binary = ET.SubElement(root, "binary")
                binary.set("id", "cover.jpg")
                content_type = "image/jpeg"
                if self.current_project.cover_path:
                    ext = os.path.splitext(self.current_project.cover_path)[1].lower()
                    if ext == '.png':
                        content_type = "image/png"
                    elif ext == '.gif':
                        content_type = "image/gif"
                    elif ext == '.bmp':
                        content_type = "image/bmp"
                binary.set("content-type", content_type)
                binary.text = self.current_project.cover_data

            book_title = ET.SubElement(title_info, "book-title")
            book_title.text = self.current_project.name

            author = ET.SubElement(title_info, "author")
            first_name = ET.SubElement(author, "first-name")
            last_name = ET.SubElement(author, "last-name")
            author_name = self.current_project.settings.get('author', 'Автор')
            if ' ' in author_name:
                parts = author_name.split(' ', 1)
                first_name.text = parts[0]
                last_name.text = parts[1] if len(parts) > 1 else ""
            else:
                first_name.text = author_name
                last_name.text = ""

            lang = ET.SubElement(title_info, "lang")
            lang.text = self.current_project.settings.get('language', 'ru')

            if self.current_project.annotation:
                annotation = ET.SubElement(title_info, "annotation")
                p = ET.SubElement(annotation, "p")
                p.text = self.current_project.annotation

            body = ET.SubElement(root, "body")

            for chapter in self.current_project.chapters:
                section = ET.SubElement(body, "section")
                title = ET.SubElement(section, "title")
                title_p = ET.SubElement(title, "p")
                title_p.text = chapter.title
                paragraphs = chapter.content.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        lines = para.strip().split('\n')
                        for line in lines:
                            if line.strip():
                                p = ET.SubElement(section, "p")
                                p.text = line.strip()

            xml_str = ET.tostring(root, encoding='utf-8', method='xml')
            dom = xml.dom.minidom.parseString(xml_str)
            pretty_xml = dom.toprettyxml(indent="  ", encoding='utf-8')

            if not pretty_xml.startswith(b'<?xml'):
                pretty_xml = b'<?xml version="1.0" encoding="utf-8"?>\n' + pretty_xml

            with open(path, 'wb') as f:
                f.write(pretty_xml)

        except Exception as e:
            print(f"Ошибка экспорта в FB2: {e}")
            import traceback
            traceback.print_exc()
            raise

    def export_txt_fallback(self, path: str):
        """Экспорт в TXT (встроенный метод)"""
        try:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(f"{self.current_project.name}\n")
                f.write(f"Автор: {self.current_project.settings.get('author', 'Автор')}\n")
                if self.current_project.cover_path:
                    f.write(f"Обложка: {os.path.basename(self.current_project.cover_path)}\n")
                f.write("=" * 60 + "\n\n")

                for chapter in self.current_project.chapters:
                    f.write(f"\n\n{chapter.title}\n")
                    f.write("-" * 40 + "\n\n")
                    f.write(chapter.content)
                    f.write("\n\n")
        except Exception as e:
            print(f"Ошибка экспорта в TXT: {e}")
            raise

    def export_pdf_fallback(self, path: str):
        """Экспорт в PDF (встроенный метод)"""
        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(path)

            doc = QTextDocument()

            html = "<html><body>"

            if self.current_project.cover_path:
                html += f'<div style="text-align: center; margin: 50px;">'
                html += f'<img src="file:///{self.current_project.cover_path}" style="max-width: 400px; max-height: 600px;">'
                html += f'</div>'
                html += '<div style="page-break-after: always;"></div>'

            html += f"<h1 style='text-align: center;'>{self.current_project.name}</h1>"
            html += f"<p style='text-align: center;'><i>{self.current_project.settings.get('author', 'Автор')}</i></p>"
            html += "<hr/>"

            for chapter in self.current_project.chapters:
                html += f"<h2>{chapter.title}</h2>"
                text = chapter.content.replace('\n', '<br>')
                html += f"<p>{text}</p>"
                html += "<div style='page-break-after: always;'></div>"

            html += "</body></html>"

            doc.setHtml(html)
            doc.print_(printer)
        except Exception as e:
            print(f"Ошибка экспорта в PDF: {e}")
            raise

    # ===== ЗАКРЫТИЕ =====

    def closeEvent(self, event):
        """При закрытии"""
        try:
            # Проверяем есть ли несохранённые изменения
            modified = any(
                isinstance(w.widget(), ChapterEditor) and
                w.widget().editor.document().isModified()
                for w in self.mdi_area.subWindowList()
            )

            if modified:
                reply = QMessageBox.question(
                    self,
                    "Подтверждение",
                    "Есть несохранённые изменения. Сохранить перед выходом?",
                    QMessageBox.StandardButton.Yes |
                    QMessageBox.StandardButton.No |
                    QMessageBox.StandardButton.Cancel
                )

                if reply == QMessageBox.StandardButton.Yes:
                    self.save_all()
                    event.accept()
                elif reply == QMessageBox.StandardButton.No:
                    event.accept()
                else:
                    event.ignore()
            else:
                event.accept()
        except Exception as e:
            print(f"Ошибка при закрытии: {e}")
            event.accept()
