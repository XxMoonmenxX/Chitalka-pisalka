# exporters/docx_exporter.py
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


class DOCXExporter:
    @staticmethod
    def export(project, path: str):
        doc = Document()

        # Заголовок
        title = doc.add_heading(project.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Автор
        author = doc.add_paragraph(f"Автор: {project.settings.get('author', 'Автор')}")
        author.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # Обложка (если есть)
        if project.cover_path and os.path.exists(project.cover_path):
            try:
                doc.add_picture(project.cover_path, width=Inches(4))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_page_break()
            except:
                pass

        # Главы
        for chapter in project.chapters:
            doc.add_heading(chapter.title, level=1)
            # Разбиваем на параграфы
            for para in chapter.content.split('\n\n'):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
            doc.add_page_break()

        doc.save(path)
